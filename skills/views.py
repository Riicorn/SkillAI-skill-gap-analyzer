from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.contrib import messages

from .models import JobRole, Skill, UserSkill
from .utils import calculate_skill_gap
import pdfplumber
from io import BytesIO
from django.http import JsonResponse
from docx import Document
# Homepage
def home(request):
    return HttpResponse("SkillAI is Running 🚀")


# Landing page
def landing(request):
    return render(request, "landing.html")


# My Skills page


@login_required
def skills_view(request):
    user_skills = UserSkill.objects.filter(user=request.user).select_related('skill')
    all_skills = Skill.objects.all().order_by("name")

    if request.method == "POST":
        skill_name_str = request.POST.get('skill_name')
        level = request.POST.get('level')

        if skill_name_str and level:
            skill_obj = Skill.objects.get(name=skill_name_str)

            UserSkill.objects.update_or_create(
                user=request.user,
                skill=skill_obj,
                defaults={'level': int(level)}
            )

            messages.success(request, f"{skill_name_str} skill updated!")
            return redirect('skills')

    return render(request, "skills/skills.html", {
        "skills": user_skills,
        "all_skills": all_skills
    })

# Skill Gap page
@login_required
def skill_gap_view(request):

    role_id = request.session.get("target_role_id")

    if role_id:
        job_role = JobRole.objects.filter(id=role_id).first()
    else:
        job_role = JobRole.objects.first()

    if not job_role:
        return render(request, "skills/skill_gap.html", {"no_roles": True})

    user_skills = UserSkill.objects.filter(user=request.user)

    if not user_skills.exists():
        return render(request, "skills/skill_gap.html", {
            "target_role": job_role,
            "has_skills": False
        })

    gaps, total_gap_score = calculate_skill_gap(request.user, job_role)

    # sort weakest skills first
    gaps = dict(sorted(gaps.items(), key=lambda x: x[1]))

    context = {
        "target_role": job_role,
        "gaps": gaps,
        "total_gap_score": total_gap_score,
        "has_skills": True
    }

    return render(request, "skills/skill_gap.html", context)

# from PyPDF2 import PdfReader
# from docx import Document
# from django.contrib.auth.decorators import login_required
# from django.shortcuts import redirect
# from django.contrib import messages
# from .models import Skill, UserSkill


# @login_required
# def upload_resume(request):

#     if request.method == "POST" and request.FILES.get("resume"):

#         resume_file = request.FILES["resume"]
#         text = ""

#         try:

#             # ---------- PDF ----------
#             if resume_file.name.endswith(".pdf"):

#                 reader = PdfReader(resume_file)

#                 for page in reader.pages:
#                     page_text = page.extract_text()
#                     if page_text:
#                         text += page_text


#             # ---------- DOCX ----------
#             elif resume_file.name.endswith(".docx"):

#                 doc = Document(resume_file)

#                 for para in doc.paragraphs:
#                     text += para.text


#             text = text.lower()

#             skills = Skill.objects.all()

#             extracted = 0

#             for skill in skills:

#                 if skill.name.lower() in text:

#                     UserSkill.objects.update_or_create(
#                         user=request.user,
#                         skill=skill,
#                         defaults={"level": 3}
#                     )

#                     extracted += 1

#             messages.success(request, f"{extracted} skills extracted from resume.")

#         except Exception as e:

#             messages.error(request, "Resume parsing failed.")

#         return redirect("skills")

#     return redirect("skills")

from PyPDF2 import PdfReader
import pdfplumber
from io import BytesIO
from docx import Document
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from .models import Skill, UserSkill


@login_required
def upload_resume(request):

    if request.method == "POST" and request.FILES.get("resume"):

        resume_file = request.FILES["resume"]
        text = ""

        try:

            # ---------- PDF ----------
            if resume_file.name.endswith(".pdf"):

                # Try PyPDF2 first
                reader = PdfReader(resume_file)

                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text

                # If PyPDF2 failed → fallback to pdfplumber
                if not text.strip():

                    resume_file.seek(0)
                    pdf_bytes = resume_file.read()

                    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text


            # ---------- DOCX ----------
            elif resume_file.name.endswith(".docx"):

                doc = Document(resume_file)

                for para in doc.paragraphs:
                    text += para.text


            text = text.lower()

            skills = Skill.objects.all()

            extracted_skills = []

            for skill in skills:

                if skill.name.lower() in text:

                    UserSkill.objects.update_or_create(
                        user=request.user,
                        skill=skill,
                        defaults={"level": 3}
                    )

                    extracted_skills.append(skill.name)

            return JsonResponse({
                "success": True,
                "skills": extracted_skills,
                "count": len(extracted_skills)
            })

        except Exception as e:

            return JsonResponse({
                "success": False,
                "error": str(e)
            })

    return JsonResponse({"success": False})

@login_required
def delete_skill(request, id):
    skill = UserSkill.objects.get(id=id, user=request.user)
    skill.delete()
    messages.success(request, "Skill removed successfully.")
    return redirect("skills")